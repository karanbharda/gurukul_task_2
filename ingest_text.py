#ingest_texts,py
import os
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings
from langchain.docstore.document import Document
import uuid
import re
import PyPDF2
from docx import Document as DocxDocument

def load_file(file_path):
    """Load content from PDF, DOCX, or TXT files."""
    file_extension = os.path.splitext(file_path)[1].lower()
    content = ""
    
    try:
        if file_extension == ".pdf":
            with open(file_path, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text = page.extract_text() or ""
                    content += text + "\n"
        
        elif file_extension == ".docx":
            doc = DocxDocument(file_path)
            for para in doc.paragraphs:
                content += para.text + "\n"
        
        elif file_extension == ".txt":
            with open(file_path, "r", encoding="utf-8") as file:
                content = file.read()
        
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
        
        return content.strip()
    
    except Exception as e:
        print(f"Error loading {file_path}: {str(e)}")
        return ""

def load_and_process_texts(file_configs):
    """Load and process texts from files with metadata."""
    documents = []
    
    # Define verse pattern for splitting and metadata
    verse_pattern = r'(?:Chapter|Verse|Mandala|Hymn|Bg|Sloka)?\s*\d+\.\d+\s*:?'
    
    # Initialize text splitter
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,  # Increased for more context
        chunk_overlap=100,
        separators=["\n\n", verse_pattern, "\n", "।।", ".", " "],
        is_separator_regex=True
    )
    
    for config in file_configs:
        file_path = config["file_path"]
        base_metadata = config["metadata"]
        
        # Load file content
        content = load_file(file_path)
        if not content:
            continue
        
        # Split text while preserving shloka structure
        chunks = text_splitter.split_text(content)
        
        # Add metadata to each chunk
        for chunk in chunks:
            # Extract verse number (expanded pattern)
            verse_match = re.search(r'(?:Bg|Verse|Mandala|Hymn|Sloka|Chapter)\s*(\d+\.\d+)', chunk, re.IGNORECASE)
            verse_number = verse_match.group(1) if verse_match else "unknown"
            
            # Attempt to extract chapter if possible
            chapter_match = re.search(r'(?:Chapter|Adhyaya)\s*(\d+)', chunk, re.IGNORECASE)
            chapter = chapter_match.group(1) if chapter_match else base_metadata.get("chapter", "unknown")
            
            # Create metadata for chunk
            chunk_metadata = {
                **base_metadata,
                "verse": verse_number,
                "chapter": chapter,
                "chunk_id": str(uuid.uuid4())
            }
            
            documents.append(Document(page_content=chunk, metadata=chunk_metadata))
    
    return documents

def create_vector_store(documents):
    """Create FAISS vector store from documents."""
    try:
        embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
        vector_store = FAISS.from_documents(documents, embeddings)
        return vector_store
    except Exception as e:
        print(f"Error creating vector store: {str(e)}")
        return None

def save_vector_store(vector_store, path="faiss_index"):
    """Save vector store locally."""
    try:
        vector_store.save_local(path)
        print(f"Vector store saved to {path}")
    except Exception as e:
        print(f"Error saving vector store: {str(e)}")

def main():
    # Define file configurations for the two PDFs
    file_configs = [
        {
            "file_path": "data/108upanishads.pdf",
            "metadata": {
                "scripture": "108 Upanishads",
                "chapter": "",
                "topic": "",
                "source": ""
            }
        },
        {
            "file_path": "data/Bhagavad-gita-As-It-Is.pdf",
            "metadata": {
                "scripture": "Bhagavad Gita",
                "chapter": "",
                "topic": "",
                "source": ""
            }
        },
        {
            "file_path": "data/Four-Vedas-English-Translation.pdf",
            "metadata": {
                "scripture": "Four Vedas",
                "chapter": "",
                "topic": "",
                "source": ""
            }
        },
        {
            "file_path": "data/ramayan.pdf",
            "metadata": {
                "scripture": "Four Vedas",
                "chapter": "",
                "topic": "",
                "source": ""
            }
        }
    ]
    
    # Load and process texts
    documents = load_and_process_texts(file_configs)
    print(f"Processed {len(documents)} document chunks")
    
    # Create and save vector store
    vector_store = create_vector_store(documents)
    if vector_store:
        save_vector_store(vector_store, path="faiss_index")
        
        # Create retriever for testing
        retriever = vector_store.as_retriever(search_kwargs={"k": 3})
        print("Retriever created. Example query usage:")
    else:
        print("Failed to create vector store.")

if __name__ == "__main__":
    main()
    