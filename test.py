#ingest_texts,py
import os
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings
from langchain.docstore.document import Document
import uuid
import re
import PyPDF2
from docx import Document as DocxDocument

def load_file(file_path):
    """Load content from PDF, DOCX, or TXT files."""
    file_extension = os.path.splitext(file_path)[1].lower()
    content = ""
    
    try:
        if file_extension == ".pdf":
            with open(file_path, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text = page.extract_text() or ""
                    content += text + "\n"
        
        elif file_extension == ".docx":
            doc = DocxDocument(file_path)
            for para in doc.paragraphs:
                content += para.text + "\n"
        
        elif file_extension == ".txt":
            with open(file_path, "r", encoding="utf-8") as file:
                content = file.read()
        
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
        
        return content.strip()
    
    except Exception as e:
        print(f"Error loading {file_path}: {str(e)}")
        return ""

def load_and_process_texts(file_configs):
    """Load and process texts from files with metadata."""
    documents = []
    
    # Define verse pattern for splitting and metadata
    verse_pattern = r'(?:Chapter|Verse|Mandala|Hymn|Bg|Sloka)?\s*\d+\.\d+\s*:?'
    
    # Initialize text splitter
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,  # Increased for more context
        chunk_overlap=100,
        separators=["\n\n", verse_pattern, "\n", "।।", ".", " "],
        is_separator_regex=True
    )
    
    for config in file_configs:
        file_path = config["file_path"]
        base_metadata = config["metadata"]
        
        # Load file content
        content = load_file(file_path)
        if not content:
            continue
        
        # Split text while preserving shloka structure
        chunks = text_splitter.split_text(content)
        
        # Add metadata to each chunk
        for chunk in chunks:
            # Extract verse number (expanded pattern)
            verse_match = re.search(r'(?:Bg|Verse|Mandala|Hymn|Sloka|Chapter)\s*(\d+\.\d+)', chunk, re.IGNORECASE)
            verse_number = verse_match.group(1) if verse_match else "unknown"
            
            # Attempt to extract chapter if possible
            chapter_match = re.search(r'(?:Chapter|Adhyaya)\s*(\d+)', chunk, re.IGNORECASE)
            chapter = chapter_match.group(1) if chapter_match else base_metadata.get("chapter", "unknown")
            
            # Create metadata for chunk
            chunk_metadata = {
                **base_metadata,
                "verse": verse_number,
                "chapter": chapter,
                "chunk_id": str(uuid.uuid4())
            }
            
            documents.append(Document(page_content=chunk, metadata=chunk_metadata))
    
    return documents

def create_vector_store(documents):
    """Create FAISS vector store from documents."""
    try:
        embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
        vector_store = FAISS.from_documents(documents, embeddings)
        return vector_store
    except Exception as e:
        print(f"Error creating vector store: {str(e)}")
        return None

def save_vector_store(vector_store, path="faiss_index"):
    """Save vector store locally."""
    try:
        vector_store.save_local(path)
        print(f"Vector store saved to {path}")
    except Exception as e:
        print(f"Error saving vector store: {str(e)}")

def main():
    # Define file configurations for the two PDFs
    file_configs = [
        {
            "file_path": "data/108upanishads.pdf",
            "metadata": {
                "scripture": "108 Upanishads",
                "chapter": "",
                "topic": "",
                "source": ""
            }
        },
        {
            "file_path": "data/Bhagavad-gita-As-It-Is.pdf",
            "metadata": {
                "scripture": "Bhagavad Gita",
                "chapter": "",
                "topic": "",
                "source": ""
            }
        },
        {
            "file_path": "data/Four-Vedas-English-Translation.pdf",
            "metadata": {
                "scripture": "Four Vedas",
                "chapter": "",
                "topic": "",
                "source": ""
            }
        },
        {
            "file_path": "data/ramayan.pdf",
            "metadata": {
                "scripture": "Four Vedas",
                "chapter": "",
                "topic": "",
                "source": ""
            }
        }
    ]
    
    # Load and process texts
    documents = load_and_process_texts(file_configs)
    print(f"Processed {len(documents)} document chunks")
    
    # Create and save vector store
    vector_store = create_vector_store(documents)
    if vector_store:
        save_vector_store(vector_store, path="faiss_index")
        
        # Create retriever for testing
        retriever = vector_store.as_retriever(search_kwargs={"k": 3})
        print("Retriever created. Example query usage:")
    else:
        print("Failed to create vector store.")

if __name__ == "__main__":
    main()
    
    
#rag_lesson.py
from fastapi import FastAPI, HTTPException
from pydantic import BaseModel
from contextlib import asynccontextmanager
from collections import Counter
import re
import logging
import pymongo
from pymongo.errors import ConnectionFailure
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings
from datetime import datetime
import uuid
import os

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# MongoDB setup
try:
    mongo_client = pymongo.MongoClient("mongodb://localhost:27017", serverSelectionTimeoutMS=5000)
    mongo_client.admin.command("ping")
    db = mongo_client["gurukul2"]
    interaction_logs = db["interaction_logs"]
    logger.info("Connected to MongoDB")
except (ConnectionFailure, pymongo.errors.ServerSelectionTimeoutError) as e:
    logger.error(f"Failed to connect to MongoDB: {str(e)}")
    interaction_logs = None

def load_vector_store(path="faiss_index"):
    """Load the FAISS vector store."""
    try:
        embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
        vector_store = FAISS.load_local(path, embeddings, allow_dangerous_deserialization=True)
        return vector_store
    except Exception as e:
        logger.error(f"Error loading vector store: {str(e)}")
        return None

def retrieve_relevant_chunks(vector_store, query, k=3):
    """Retrieve top k relevant chunks for the query."""
    retriever = vector_store.as_retriever(search_kwargs={"k": k})
    try:
        docs = retriever.invoke(query)
        return docs
    except Exception as e:
        logger.error(f"Error retrieving chunks: {str(e)}")
        return []

def extract_verses_from_content(content):
    """Extract verse numbers from chunk content as a fallback."""
    verse_pattern = r'(?:Bg|Verse|Mandala|Hymn|Sloka|[A-Za-z\s]*?)\s*(\d+\.\d+[\s]*(?:[-–]\d+)?(?:\s*[a-zA-Z]?)?)'
    matches = re.findall(verse_pattern, content, re.IGNORECASE)
    return matches if matches else ["unknown"]

def format_lesson(query, documents):
    """Format the retrieved documents into a structured lesson."""
    if not documents:
        logger.warning("No documents retrieved for query")
        return {
            "theme": "No Relevant Information Found",
            "references": ["None"],
            "explanation": "No relevant content was found for the query. Try rephrasing or exploring a related topic.",
            "activity": "Consider researching the topic further or asking a more specific question about the scriptures."
        }

    # Clean and generate a descriptive theme
    cleaned_query = re.sub(r'\b(about|is|what|the)\b', '', query, flags=re.IGNORECASE).strip()
    theme = cleaned_query.capitalize()
    if any(word in query.lower() for word in ["what", "nature"]):
        theme = f"The Essence of {cleaned_query.capitalize()}"
    elif "essence of" in query.lower():
        theme = f"Understanding {cleaned_query.split('essence of')[-1].strip().capitalize()}"

    # Collect reference verses
    references = []
    for doc in documents:
        verse = doc.metadata.get("verse", "unknown")
        scripture = doc.metadata.get("scripture", "Unknown Scripture")
        chapter = doc.metadata.get("chapter", "unknown")
        if verse == "unknown":
            content_verses = extract_verses_from_content(doc.page_content)
            for v in content_verses:
                if v != "unknown":
                    ref = f"{scripture} {v}"
                    if chapter != "unknown":
                        ref += f" (Chapter {chapter})"
                    references.append(ref)
            if content_verses == ["unknown"]:
                logger.debug(f"No verses found in chunk (scripture: {scripture}): {doc.page_content[:200]}")
        else:
            ref = f"{scripture} {verse}"
            if chapter != "unknown":
                ref += f" (Chapter {chapter})"
            references.append(ref)
    references = list(set(references))
    if not references or all("unknown" in ref.lower() for ref in references):
        references = ["No specific verses identified"]

    # Determine primary scripture dynamically (prefer Bhagavad Gita if present)
    scriptures = [doc.metadata.get("scripture", "scriptures") for doc in documents]
    scripture_counts = Counter(scriptures)
    primary_scripture = "Bhagavad Gita" if "Bhagavad Gita" in scripture_counts else max(scripture_counts, key=scripture_counts.get) if scripture_counts else "scriptures"
    scriptures_list = sorted(scripture_counts, key=scripture_counts.get, reverse=True)
    scriptures_text = ", ".join(scriptures_list) if len(scriptures_list) > 1 else scriptures_list[0] if scriptures_list else "scriptures"

    # Generate a cohesive explanation
    explanation = f"The {scriptures_text} provide profound wisdom on {theme.lower()}. "
    key_points = []
    for doc in documents:
        content = doc.page_content.replace("\n", " ").strip()
        sentences = [s.strip() for s in content.split(". ") if s.strip() and len(s) > 10]
        first_sentence = sentences[0] + "." if sentences else content[:200] + "..."
        if first_sentence[0].islower() or first_sentence.startswith("f "):
            first_sentence = re.sub(r'^\w\s+', '', first_sentence).capitalize()
        key_points.append(first_sentence)
    if key_points:
        explanation += "Key insights include: " + "; ".join(f"{i+1}. {point}" for i, point in enumerate(key_points)) + " "
    explanation += f"Together, these teachings emphasize that {theme.lower()} involves aligning one’s life with the eternal principles of existence, fostering inner peace and purpose."

    # Generate a specific activity/story prompt
    character_context = "Arjuna" if primary_scripture == "Bhagavad Gita" else "a Vedic sage" if primary_scripture == "108 Upanishads" else "a seeker"
    activity = (
        f"Envision yourself as {character_context} studying the {primary_scripture}, contemplating {theme.lower()}. "
        f"Write a short story or dialogue where you face a challenge related to {theme.lower()} "
        f"(e.g., seeking life’s purpose, resolving a moral dilemma, or overcoming doubt). "
        f"How would you draw upon the insights above to find clarity or purpose? "
        f"Alternatively, reflect in a journal entry on how these teachings could illuminate a personal experience."
    )

    return {
        "theme": theme,
        "references": references,
        "explanation": explanation,
        "activity": activity
    }

def rag_lesson_pipeline(query):
    """Run the RAG pipeline to generate a lesson from a query."""
    vector_store = load_vector_store()
    if not vector_store:
        return None
    documents = retrieve_relevant_chunks(vector_store, query)
    lesson = format_lesson(query, documents)
    return lesson

# Lifespan handler for startup and shutdown
@asynccontextmanager
async def lifespan(app: FastAPI):
    vector_store = load_vector_store()
    if not vector_store:
        logger.error("Failed to load vector store. Ensure faiss_index/ exists. Run ingest_text.py to create it.")
    else:
        logger.info("RAG pipeline initialized successfully")
    
    if interaction_logs is None:
        logger.error("MongoDB connection not established")
    
    yield
    
    if mongo_client is not None:
        mongo_client.close()
        logger.info("Closed MongoDB connection")

# Initialize FastAPI app
app = FastAPI(
    title="task2",
    description="API for querying Vedic texts with RAG pipeline",
    lifespan=lifespan
)

# Pydantic model for request
class QueryRequest(BaseModel):
    query: str

# Pydantic model for response
class LessonResponse(BaseModel):
    theme: str
    references: list[str]
    explanation: str
    activity: str

@app.post("/ask-vedas", response_model=LessonResponse)
async def ask_vedas(request: QueryRequest):
    """Process a query and return a lesson-formatted response."""
    if not request.query.strip():
        logger.warning("Empty query received")
        raise HTTPException(status_code=400, detail="Query cannot be empty")

    # Generate a unique ID for logging
    user_id = str(uuid.uuid4())
    
    vector_store = load_vector_store()
    if not vector_store:
        logger.error("Vector store not found")
        raise HTTPException(status_code=500, detail="Vector store not available")

    try:
        lesson = rag_lesson_pipeline(request.query)
        if not lesson:
            logger.error("RAG pipeline failed to generate lesson")
            raise HTTPException(status_code=500, detail="Failed to generate lesson")
    except Exception as e:
        logger.error(f"RAG pipeline error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"RAG pipeline error: {str(e)}")

    # Retrieve documents for logging
    try:
        docs = vector_store.as_retriever(search_kwargs={"k": 3}).invoke(request.query)
        retrieved_docs = [
            {
                "content": doc.page_content[:200] + ("..." if len(doc.page_content) > 200 else ""),
                "metadata": doc.metadata
            }
            for doc in docs
        ]
    except Exception as e:
        logger.warning(f"Failed to retrieve documents for logging: {str(e)}")
        retrieved_docs = []

    # Log interaction to MongoDB
    if interaction_logs is not None:
        try:
            log_entry = {
                "user_id": user_id,
                "query": request.query,
                "timestamp": datetime.utcnow(),
                "retrieved_docs": retrieved_docs,
                "response": lesson
            }
            interaction_logs.insert_one(log_entry)
            logger.info(f"Logged interaction for user_id: {user_id}")
        except (ConnectionFailure, pymongo.errors.PyMongoError) as e:
            logger.error(f"Failed to log to MongoDB: {str(e)}")

    return lesson

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)